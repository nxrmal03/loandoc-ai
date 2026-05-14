"""
Document ingestion pipeline:
  1. Save uploaded file to disk
  2. Extract text (PDF via pypdf, DOCX via python-docx)
  3. Chunk with RecursiveCharacterTextSplitter
  4. Embed via OpenAI text-embedding-3-small
  5. Store in persistent local Chroma collection
"""
import json
import logging
import uuid
from datetime import datetime
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings

from app.core.config import settings
from app.models.schemas import DocumentMeta

logger = logging.getLogger(__name__)

METADATA_FILE = "documents_meta.json"


class DocumentProcessor:
    def __init__(self) -> None:
        self._upload_dir = settings.upload_dir_abs
        self._chroma_path = settings.chroma_db_path_abs
        self._meta_path = self._upload_dir.parent / METADATA_FILE

        self._upload_dir.mkdir(parents=True, exist_ok=True)
        self._chroma_path.mkdir(parents=True, exist_ok=True)

        self._embeddings = GoogleGenerativeAIEmbeddings(
            model=settings.embedding_model,
            google_api_key=settings.google_api_key,
        )
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        self._vectorstore = Chroma(
            collection_name="loandoc_chunks",
            embedding_function=self._embeddings,
            persist_directory=str(self._chroma_path),
        )

    # ── Public API ─────────────────────────────────────────────────────────────

    async def process_upload(
        self, filename: str, content: bytes, content_type: str
    ) -> dict:
        """Save, parse, chunk, embed, and index an uploaded document."""
        document_id = str(uuid.uuid4())
        suffix = Path(filename).suffix.lower()
        dest = self._upload_dir / f"{document_id}{suffix}"
        dest.write_bytes(content)
        logger.info("Saved %s → %s", filename, dest)

        try:
            raw_docs = self._extract(dest, suffix, document_id, filename)
        except Exception as exc:
            dest.unlink(missing_ok=True)
            raise ValueError(f"Text extraction failed for '{filename}': {exc}") from exc

        chunks = self._splitter.split_documents(raw_docs)
        if not chunks:
            dest.unlink(missing_ok=True)
            raise ValueError(f"No extractable text found in '{filename}'.")

        # Tag every chunk with a stable chunk_id
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"] = f"{document_id}_chunk_{i}"
            chunk.metadata["document_id"] = document_id

        self._vectorstore.add_documents(chunks)
        logger.info("Indexed %d chunks for document %s", len(chunks), document_id)

        meta = DocumentMeta(
            document_id=document_id,
            filename=filename,
            file_size_bytes=len(content),
            chunk_count=len(chunks),
            uploaded_at=datetime.utcnow(),
            content_type=content_type,
        )
        self._save_meta(meta)

        return {
            "document_id": document_id,
            "filename": filename,
            "chunk_count": len(chunks),
        }

    def list_documents(self) -> list[DocumentMeta]:
        """Return metadata for all ingested documents."""
        return [DocumentMeta(**m) for m in self._load_all_meta().values()]

    def delete_document(self, document_id: str) -> bool:
        """Remove document from disk and Chroma. Returns False if not found."""
        all_meta = self._load_all_meta()
        if document_id not in all_meta:
            return False

        meta = all_meta[document_id]
        suffix = Path(meta["filename"]).suffix.lower()
        dest = self._upload_dir / f"{document_id}{suffix}"
        dest.unlink(missing_ok=True)

        # Delete all chunks belonging to this document from Chroma
        results = self._vectorstore.get(where={"document_id": document_id})
        if results and results.get("ids"):
            self._vectorstore.delete(ids=results["ids"])

        del all_meta[document_id]
        self._meta_path.write_text(json.dumps(all_meta, default=str), encoding="utf-8")
        logger.info("Deleted document %s", document_id)
        return True

    def get_vectorstore(self) -> Chroma:
        return self._vectorstore

    # ── Text extraction ────────────────────────────────────────────────────────

    def _extract(
        self, path: Path, suffix: str, document_id: str, filename: str
    ) -> list[Document]:
        if suffix == ".pdf":
            return self._extract_pdf(path, document_id, filename)
        if suffix == ".docx":
            return self._extract_docx(path, document_id, filename)
        raise ValueError(f"Unsupported file type: {suffix}. Only PDF and DOCX are accepted.")

    def _extract_pdf(
        self, path: Path, document_id: str, filename: str
    ) -> list[Document]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        docs: list[Document] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "document_id": document_id,
                        "filename": filename,
                        "page_num": page_num,
                        "source": filename,
                    },
                )
            )
        return docs

    def _extract_docx(
        self, path: Path, document_id: str, filename: str
    ) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Group paragraphs into ~page-sized blocks to preserve metadata granularity
        block_size = 30
        docs: list[Document] = []
        for block_idx in range(0, len(paragraphs), block_size):
            block = "\n".join(paragraphs[block_idx : block_idx + block_size])
            docs.append(
                Document(
                    page_content=block,
                    metadata={
                        "document_id": document_id,
                        "filename": filename,
                        "page_num": (block_idx // block_size) + 1,
                        "source": filename,
                    },
                )
            )
        return docs

    # ── Metadata persistence ───────────────────────────────────────────────────

    def _load_all_meta(self) -> dict:
        if not self._meta_path.exists():
            return {}
        return json.loads(self._meta_path.read_text(encoding="utf-8"))

    def _save_meta(self, meta: DocumentMeta) -> None:
        all_meta = self._load_all_meta()
        all_meta[meta.document_id] = meta.model_dump(mode="json")
        self._meta_path.write_text(json.dumps(all_meta, default=str), encoding="utf-8")
